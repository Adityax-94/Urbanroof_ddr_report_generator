import argparse
import base64
import io
import json
import os
import sys
import time
from pathlib import Path

try:
    import fitz
except ImportError:
    sys.exit("Run: pip install pymupdf")

try:
    from google import genai
    from google.genai import types
except ImportError:
    sys.exit("Run: pip install google-genai")

try:
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("Run: pip install python-docx")


def extract_text(pdf_path):
    doc = fitz.open(pdf_path)
    pages = []
    for i, page in enumerate(doc, 1):
        text = page.get_text("text").strip()
        if text:
            pages.append(f"[PAGE {i}]\n{text}")
    doc.close()
    return "\n\n".join(pages)


def extract_images(pdf_path, max_images=5, min_w=200, min_h=200):
    doc = fitz.open(pdf_path)
    candidates = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        for idx, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_img = doc.extract_image(xref)
            w, h = base_img["width"], base_img["height"]
            if w < min_w or h < min_h:
                continue
            candidates.append({
                "page": page_num + 1,
                "index": idx,
                "bytes": base_img["image"],
                "ext": base_img["ext"],
                "width": w,
                "height": h,
                "size": w * h,
                "source": Path(pdf_path).stem,
            })
    doc.close()
    candidates.sort(key=lambda x: x["size"], reverse=True)
    kept = candidates[:max_images]
    print(f"     {Path(pdf_path).name}: {len(candidates)} images found, using top {len(kept)}")
    return kept


PROMPT = """
You are a building diagnostics specialist. Analyze the inspection report and thermal report 
provided (text + images) and return ONLY a valid JSON object — no markdown fences, no extra text.

Rules:
- Never invent facts. Only use what is in the documents.
- Write "Not Available" for missing data.
- If data conflicts between documents, describe the conflict.
- Use simple, client-friendly language.

Return ONLY this JSON structure:
{
  "property_name": "string or Not Available",
  "report_date": "string or Not Available",
  "issue_summary": "3-5 sentence executive summary of main problems",
  "areas": [
    {
      "name": "area name",
      "inspection_finding": "what inspector observed",
      "thermal_finding": "thermal findings or Not Available",
      "image_refs": ["e.g. Page 2 Thermal Report shows 52C hotspot"],
      "combined_assessment": "merged finding from both reports",
      "severity": "Critical | High | Medium | Low",
      "severity_reason": "one sentence justification"
    }
  ],
  "root_causes": [
    { "issue": "issue name", "cause": "probable explanation" }
  ],
  "recommended_actions": [
    { "action": "what to do", "timeline": "Immediate | Within 30 days | Planned" }
  ],
  "additional_notes": "other observations or Not Available",
  "missing_information": [
    { "item": "expected item", "status": "Not Available | Conflicting: description" }
  ]
}
"""

def call_gemini(inspection_text, thermal_text, all_images, api_key):
    client = genai.Client(api_key=api_key)

    contents = []
    contents.append(
        f"=== INSPECTION REPORT ===\n{inspection_text}\n\n"
        f"=== THERMAL REPORT ===\n{thermal_text}\n\n"
        f"{PROMPT}"
    )
    for img in all_images:
        mime = "image/jpeg" if img["ext"] in ("jpg", "jpeg") else f"image/{img['ext']}"
        contents.append(
            types.Part.from_bytes(data=img["bytes"], mime_type=mime)
        )
        contents.append(f"[Image from {img['source']}, page {img['page']}, {img['width']}x{img['height']}px]")

    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=contents,
    )

    raw = response.text.strip()
    if "```" in raw:
        parts = raw.split("```")
        for part in parts:
            part = part.strip()
            if part.startswith("json"):
                part = part[4:].strip()
            try:
                return json.loads(part)
            except Exception:
                continue

    return json.loads(raw)


SEVERITY_COLORS = {
    "Critical": RGBColor(0xC0, 0x00, 0x00),
    "High":     RGBColor(0xFF, 0x66, 0x00),
    "Medium":   RGBColor(0xFF, 0xC0, 0x00),
    "Low":      RGBColor(0x00, 0x70, 0xC0),
}

def build_docx(ddr, all_images, out_path):
    doc = Document()

    t = doc.add_heading("Detailed Diagnostic Report (DDR)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Property: {ddr.get('property_name','Not Available')}  |  "
        f"Date: {ddr.get('report_date','Not Available')}  |  "
        f"Prepared by: AI Diagnostic System"
    )
    doc.add_paragraph()
    doc.add_heading("1. Property Issue Summary", 1)
    doc.add_paragraph(ddr.get("issue_summary", "Not Available"))
    doc.add_heading("2. Area-wise Observations", 1)
    for area in ddr.get("areas", []):
        doc.add_heading(area.get("name", "Unknown Area"), 2)

        for label, key in [
            ("Inspection finding",  "inspection_finding"),
            ("Thermal finding",     "thermal_finding"),
            ("Combined assessment", "combined_assessment"),
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(area.get(key, "Not Available"))

        sev = area.get("severity", "Unknown")
        p = doc.add_paragraph()
        p.add_run("Severity: ").bold = True
        run = p.add_run(f"{sev}  ")
        run.bold = True
        run.font.color.rgb = SEVERITY_COLORS.get(sev, RGBColor(0x40, 0x40, 0x40))
        p.add_run(f"— {area.get('severity_reason', '')}")

        refs = area.get("image_refs", [])
        if refs:
            p = doc.add_paragraph()
            p.add_run("Image references: ").bold = True
            p.add_run(" | ".join(refs))

        embedded = False
        for img in all_images:
            ref_combined = " ".join(refs).lower()
            if f"page {img['page']}" in ref_combined or img["source"].lower() in ref_combined:
                try:
                    doc.add_picture(io.BytesIO(img["bytes"]), width=Inches(4.5))
                    cap = doc.add_paragraph(f"Fig: {img['source']} — page {img['page']}")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    embedded = True
                except Exception:
                    pass
                break

        if refs and not embedded:
            doc.add_paragraph("Image Not Available")

        doc.add_paragraph()
    doc.add_heading("3. Probable Root Cause", 1)
    for rc in ddr.get("root_causes", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{rc.get('issue','')}: ").bold = True
        p.add_run(rc.get("cause", "Not Available"))

    doc.add_heading("4. Severity Assessment", 1)
    areas = ddr.get("areas", [])
    if areas:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Area", "Severity", "Reasoning"
        for area in areas:
            row = table.add_row().cells
            row[0].text = area.get("name", "")
            row[1].text = area.get("severity", "")
            row[2].text = area.get("severity_reason", "")
    else:
        doc.add_paragraph("Not Available")
    doc.add_heading("5. Recommended Actions", 1)
    for action in ddr.get("recommended_actions", []):
        p = doc.add_paragraph(style="List Number")
        p.add_run(action.get("action", ""))
        p.add_run(f"  [{action.get('timeline','')}]").italic = True
    doc.add_heading("6. Additional Notes", 1)
    doc.add_paragraph(ddr.get("additional_notes", "Not Available"))
    doc.add_heading("7. Missing or Unclear Information", 1)
    missing = ddr.get("missing_information", [])
    if missing:
        for item in missing:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{item.get('item','')}: ").bold = True
            p.add_run(item.get("status", "Not Available"))
    else:
        doc.add_paragraph("All expected information was present in the documents.")

    doc.save(out_path)
    print(f"✅  DDR saved to: {out_path}")

def main():
    parser = argparse.ArgumentParser(description="DDR Report Generator (Gemini)")
    parser.add_argument("--inspection", required=True)
    parser.add_argument("--thermal",    required=True)
    parser.add_argument("--out",        default="DDR_Report.docx")
    parser.add_argument("--api-key",    default=os.environ.get("GEMINI_API_KEY", ""))
    args = parser.parse_args()

    if not args.api_key:
        sys.exit("Provide your Gemini key via --api-key YOUR_KEY\nGet a free key at: https://aistudio.google.com")

    print("📄  Extracting text...")
    inspection_text = extract_text(args.inspection)
    thermal_text    = extract_text(args.thermal)

    print("🖼   Extracting images (top 5 per document)...")
    inspection_images = extract_images(args.inspection, max_images=5)
    thermal_images    = extract_images(args.thermal,    max_images=5)
    all_images = inspection_images + thermal_images
    print(f"     Sending {len(all_images)} images total to Gemini")

    print("🤖  Calling Gemini AI...")
    ddr = call_gemini(inspection_text, thermal_text, all_images, args.api_key)

    print("📝  Building Word document...")
    build_docx(ddr, all_images, args.out)

if __name__ == "__main__":
    main()
