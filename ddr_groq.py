"""
DDR Report Generator v2 — Groq (free, fast, works in India)
=============================================================
1. Get free key at: https://console.groq.com  (Sign up → API Keys → Create)
2. pip install groq pymupdf python-docx
3. python ddr_groq.py --inspection inspection.pdf --thermal thermal.pdf --api-key gsk_...
"""

import argparse
import io
import json
import os
import sys
from pathlib import Path

try:
    import fitz  # PyMuPDF
except ImportError:
    sys.exit("Run: pip install pymupdf")

try:
    from groq import Groq
except ImportError:
    sys.exit("Run: pip install groq")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("Run: pip install python-docx")


# ── PDF helpers ───────────────────────────────────────────────────────────────

def extract_text(pdf_path):
    doc = fitz.open(pdf_path)
    pages = []
    for i, page in enumerate(doc, 1):
        text = page.get_text("text").strip()
        if text:
            pages.append(f"[PAGE {i}]\n{text}")
    doc.close()
    full = "\n\n".join(pages)
    # Groq context limit — use 18000 chars per doc (llama-3.3-70b has 128k context)
    return full[:18000]


def extract_images(pdf_path, max_images=6, min_w=150, min_h=150):
    """Extract largest images — for embedding in Word doc."""
    doc = fitz.open(pdf_path)
    candidates = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        for idx, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_img = doc.extract_image(xref)
            w, h = base_img["width"], base_img["height"]
            if w < min_w or h < min_h:
                continue
            candidates.append({
                "page": page_num + 1,
                "index": idx,
                "bytes": base_img["image"],
                "ext": base_img["ext"],
                "width": w,
                "height": h,
                "size": w * h,
                "source": Path(pdf_path).stem,
            })
    doc.close()
    candidates.sort(key=lambda x: x["size"], reverse=True)
    kept = candidates[:max_images]
    print(f"     {Path(pdf_path).name}: {len(candidates)} images found, keeping top {len(kept)}")
    return kept


# ── AI call ───────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are a senior building diagnostics specialist.
Your job is to read two documents — an Inspection Report and a Thermal Report — and produce a structured DDR (Detailed Diagnostic Report).

STRICT RULES:
1. For property_name: search the document carefully for any flat number, unit number, apartment number, address, or property identifier. Even partial info like "Flat No. 103" counts.
2. For thermal findings per area: the thermal report may not name rooms explicitly. Instead look for:
   - Any temperature readings (e.g. 28.5C, hotspot, coldspot)
   - Any anomalies described near that area's photos or page numbers
   - Any moisture or heat patterns that correspond to that area's issues
   If ANY thermal data exists in the report — even general observations — assign it to the most relevant area rather than writing "Not Available". Only write "Not Available" if the thermal report has zero data whatsoever for that area.
3. Extract ALL areas mentioned in the inspection report — do not drop any rooms or locations.
4. Never invent facts. If thermal data exists but cannot be matched to a specific room, note it under the closest area or under Additional Notes.
5. Use simple, professional, client-friendly language.
6. Recommended actions must be specific — mention exact area and issue.
7. Generate at least 6 recommended actions in priority order.
8. Return ONLY raw JSON — no markdown fences, no explanation, no extra text."""

JSON_SCHEMA = """{
  "property_name": "extracted from documents or Not Available",
  "report_date": "extracted from documents or Not Available",
  "issue_summary": "Write 4-6 complete sentences summarizing ALL major problems found across both documents. Mention specific areas affected, types of issues, and overall severity.",
  "areas": [
    {
      "name": "exact area name from documents",
      "inspection_finding": "detailed description of what the inspector observed in this area",
      "thermal_finding": "specific thermal data for this area including temperatures if available, or Not Available",
      "image_refs": ["specific image or photo reference from the documents, e.g. Photo 1-6 from Inspection Report page 1"],
      "combined_assessment": "2-3 sentences merging inspection and thermal findings into one clear assessment",
      "severity": "Critical | High | Medium | Low",
      "severity_reason": "specific one-sentence reason based on evidence from the documents"
    }
  ],
  "root_causes": [
    { "issue": "specific issue name", "cause": "detailed probable explanation based on evidence in documents" }
  ],
  "recommended_actions": [
    { "action": "specific action mentioning exact area and issue", "timeline": "Immediate | Within 30 days | Planned" }
  ],
  "additional_notes": "any patterns, recurring issues, positive findings, or observations that do not fit above sections",
  "missing_information": [
    { "item": "name of expected item", "status": "Not Available | Conflicting: exact description of conflict" }
  ]
}"""


def call_groq(inspection_text, thermal_text, api_key):
    client = Groq(api_key=api_key)

    user_message = (
        f"=== INSPECTION REPORT ===\n{inspection_text}\n\n"
        f"=== THERMAL REPORT ===\n{thermal_text}\n\n"
        f"Analyze both documents thoroughly. Pay special attention to matching "
        f"thermal findings to specific rooms from the inspection report.\n\n"
        f"Return ONLY a JSON object matching exactly this schema:\n{JSON_SCHEMA}"
    )

    print("     Sending to llama-3.3-70b...")
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": user_message}
        ],
        max_tokens=6000,   # increased from 4000
        temperature=0.1,
    )

    raw = response.choices[0].message.content.strip()

    # Strip markdown fences if model added them
    if "```" in raw:
        for part in raw.split("```"):
            part = part.strip()
            if part.startswith("json"):
                part = part[4:].strip()
            try:
                return json.loads(part)
            except Exception:
                continue

    return json.loads(raw)


# ── Word document builder ─────────────────────────────────────────────────────

SEVERITY_COLORS = {
    "Critical": RGBColor(0xC0, 0x00, 0x00),
    "High":     RGBColor(0xFF, 0x66, 0x00),
    "Medium":   RGBColor(0xFF, 0xC0, 0x00),
    "Low":      RGBColor(0x00, 0x70, 0xC0),
}

SEVERITY_BG = {
    "Critical": RGBColor(0xFF, 0xE0, 0xE0),
    "High":     RGBColor(0xFF, 0xF0, 0xE0),
    "Medium":   RGBColor(0xFF, 0xFD, 0xE0),
    "Low":      RGBColor(0xE0, 0xF0, 0xFF),
}


def set_cell_bg(cell, rgb):
    """Set background color of a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def build_docx(ddr, inspection_images, thermal_images, out_path):
    doc = Document()

    # ── Page margins (slightly narrower for more content space)
    from docx.shared import Inches as In
    for section in doc.sections:
        section.top_margin    = In(1)
        section.bottom_margin = In(1)
        section.left_margin   = In(1.2)
        section.right_margin  = In(1.2)

    # ── Title block
    t = doc.add_heading("Detailed Diagnostic Report (DDR)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Property: {ddr.get('property_name','Not Available')}  |  "
        f"Date: {ddr.get('report_date','Not Available')}  |  "
        f"Prepared by: AI Diagnostic System"
    ).italic = True
    doc.add_paragraph()

    # ── 1. Property Issue Summary
    doc.add_heading("1. Property Issue Summary", 1)
    summary = ddr.get("issue_summary", "Not Available")
    doc.add_paragraph(summary)
    doc.add_paragraph()

    # ── 2. Area-wise Observations
    doc.add_heading("2. Area-wise Observations", 1)
    areas = ddr.get("areas", [])

    # Build image pool: thermal images first (more diagnostic value), then inspection
    image_pool = thermal_images + inspection_images

    for area_idx, area in enumerate(areas):
        doc.add_heading(area.get("name", "Unknown Area"), 2)

        # Findings table for clean layout
        tbl = doc.add_table(rows=4, cols=2)
        tbl.style = "Light Grid"
        tbl.columns[0].width = Inches(1.8)
        tbl.columns[1].width = Inches(4.5)

        rows_data = [
            ("Inspection Finding", area.get("inspection_finding", "Not Available")),
            ("Thermal Finding",    area.get("thermal_finding",    "Not Available")),
            ("Combined Assessment",area.get("combined_assessment","Not Available")),
            ("Image References",  " | ".join(area.get("image_refs", ["Not Available"]))),
        ]
        for i, (label, value) in enumerate(rows_data):
            row = tbl.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value

        doc.add_paragraph()

        # Severity badge paragraph
        sev = area.get("severity", "Unknown")
        p = doc.add_paragraph()
        p.add_run("Severity: ").bold = True
        run = p.add_run(f"  {sev}  ")
        run.bold = True
        run.font.color.rgb = SEVERITY_COLORS.get(sev, RGBColor(0x40, 0x40, 0x40))
        run.font.size = Pt(12)
        p.add_run(f"   {area.get('severity_reason', '')}")

        # Embed image — pick from thermal pool first, cycle through
        if image_pool:
            img = image_pool[area_idx % len(image_pool)]
            try:
                doc.add_paragraph()
                doc.add_picture(io.BytesIO(img["bytes"]), width=Inches(4.0))
                cap = doc.add_paragraph(f"Figure {area_idx+1}: {img['source'].replace('-',' ').title()} — Page {img['page']}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
            except Exception:
                doc.add_paragraph("[ Image Not Available ]")

        doc.add_paragraph()

    # ── 3. Probable Root Cause
    doc.add_heading("3. Probable Root Cause", 1)
    for rc in ddr.get("root_causes", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{rc.get('issue','')}: ").bold = True
        p.add_run(rc.get("cause", "Not Available"))
    doc.add_paragraph()

    # ── 4. Severity Assessment table
    doc.add_heading("4. Severity Assessment", 1)
    if areas:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        # Header row
        hdr = table.rows[0].cells
        for cell, label in zip(hdr, ["Area", "Severity", "Reasoning"]):
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
        # Data rows
        for area in areas:
            row = table.add_row().cells
            sev = area.get("severity", "")
            row[0].text = area.get("name", "")
            row[1].text = sev
            row[2].text = area.get("severity_reason", "")
            # Color the severity cell
            if sev in SEVERITY_BG:
                set_cell_bg(row[1], SEVERITY_BG[sev])
                if row[1].paragraphs[0].runs:
                    row[1].paragraphs[0].runs[0].font.color.rgb = SEVERITY_COLORS[sev]
                    row[1].paragraphs[0].runs[0].bold = True
    else:
        doc.add_paragraph("Not Available")
    doc.add_paragraph()

    # ── 5. Recommended Actions
    doc.add_heading("5. Recommended Actions", 1)
    for action in ddr.get("recommended_actions", []):
        p = doc.add_paragraph(style="List Number")
        p.add_run(action.get("action", "")).bold = False
        timeline = action.get("timeline", "")
        t_run = p.add_run(f"  [{timeline}]")
        t_run.italic = True
        t_run.font.color.rgb = SEVERITY_COLORS.get(
            "Immediate" if timeline == "Immediate" else
            "High"      if timeline == "Within 30 days" else "Low",
            RGBColor(0x40, 0x40, 0x40)
        )
    doc.add_paragraph()

    # ── 6. Additional Notes
    doc.add_heading("6. Additional Notes", 1)
    doc.add_paragraph(ddr.get("additional_notes", "Not Available"))
    doc.add_paragraph()

    # ── 7. Missing or Unclear Information
    doc.add_heading("7. Missing or Unclear Information", 1)
    missing = ddr.get("missing_information", [])
    if missing:
        for item in missing:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{item.get('item','')}: ").bold = True
            p.add_run(item.get("status", "Not Available"))
    else:
        doc.add_paragraph("All expected information was present in the documents.")

    doc.save(out_path)
    print(f"✅  DDR saved to: {out_path}")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="DDR Report Generator v2 (Groq)")
    parser.add_argument("--inspection", required=True)
    parser.add_argument("--thermal",    required=True)
    parser.add_argument("--out",        default="DDR_Report.docx")
    parser.add_argument("--api-key",    default=os.environ.get("GROQ_API_KEY", ""))
    args = parser.parse_args()

    if not args.api_key:
        sys.exit(
            "Provide your Groq key via --api-key YOUR_KEY\n"
            "Get a free key at: https://console.groq.com"
        )

    print("📄  Extracting text...")
    inspection_text = extract_text(args.inspection)
    thermal_text    = extract_text(args.thermal)
    print(f"     Inspection: {len(inspection_text)} chars | Thermal: {len(thermal_text)} chars")

    print("🖼   Extracting images...")
    inspection_images = extract_images(args.inspection, max_images=6)
    thermal_images    = extract_images(args.thermal,    max_images=6)

    print("🤖  Calling Groq AI (llama-3.3-70b)...")
    ddr = call_groq(inspection_text, thermal_text, args.api_key)

    print("📝  Building Word document...")
    build_docx(ddr, inspection_images, thermal_images, args.out)

if __name__ == "__main__":
    main()
